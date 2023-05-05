from django.contrib import admin
from .models import GradesFiles, AnalysisTemplate
from .forms import FileFieldForm
from django.utils.encoding import escape_uri_path
from user.models import UserInfo

import os, re, shutil
from django.shortcuts import render, redirect
from django.http import FileResponse
import pandas as pd
import numpy as np
import matplotlib as mpl
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt,Cm    #字号
from docx.oxml.ns import qn    #字体
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING   #对齐
from docx.enum.table import WD_TABLE_ALIGNMENT    #对齐

# Register your models here.
@admin.register(AnalysisTemplate)
class AnalysisTemplateAdmin(admin.ModelAdmin):
    list_display = ['id', 'name']
    readonly_fields = ['instruction']
    fieldsets = (
        ('简介', {'fields':['name', 'instruction']}),
        ('成绩分析单', {'fields':['comment']}),
        ('教学质量分析报告', {'fields':['basic_condition', 'teaching_condition', 'experience', 'advice']}),
        ('所属教师', {'fields':['instructor']})
    )
    raw_id_fields = ('instructor',)
    search_fields = ('template_id',)
    def get_queryset(self, request):
        qs = super(AnalysisTemplateAdmin, self).get_queryset(request)
        if request.user.is_superuser:
            return qs
        return qs.filter(instructor__username=request.user)

    def get_changeform_initial_data(self, request):
        return {'instructor': request.user}

class GradesFilesAdmin(admin.ModelAdmin):
    form = FileFieldForm

    add_form_template = 'admin/grades/GradesFiles/add/add_grades_file.html'
    autocomplete_fields  = ('template_id',)

    def response_add(self, request, obj):
        if '_save' in request.POST:
            files = request.FILES.getlist('multiple_grades_files')
            f_path = ''
            # all_files_df 用以保存上传所有文件的关键信息
            all_files_df = pd.DataFrame(
                index=['课程名称', '上课教师', '班级名称', '授课时间', '课程类别', '学时', '应考人数', '缺考人数',
                       '100-90分', '89-80分', '79-70分', '69-60分', '59分及其以下', '平均分', '图片名'])
            # try:
            for file in files:
                grade_data, file_info_list, cal_method = preprocess(file)
                f_path = grade_data['上课教师'].iloc[0] + " " + xuenian_to_zirannian(
                    grade_data['开课学期'].iloc[0]) + " 期末考试归档资料" + pd.Timestamp.now().strftime('%Y%m%d%H%M')
                if not os.path.exists(f'media/grades_files/{f_path}'):
                    os.mkdir('media/grades_files/' + f_path)
                # 用以单独保存相应的图片
                if not os.path.exists(f'media/grades_files/{f_path}pic'):
                    os.mkdir('media/grades_files/' + f_path + 'pic')
                with open(f'media/grades_files/{f_path}/{file.name}', 'wb') as f:
                    for chunk in file.chunks():
                        f.write(chunk)
                course_class, file_info_l = generate_cjfxd(grade_data, file_info_list, cal_method, f_path, obj.template_id)
                # course_class必须唯一，否则下面这个语句会造成覆盖
                all_files_df[course_class] = file_info_l
            # except Exception:
            #     return redirect('/grades/upload_error')

            for i in all_files_df.T.groupby('课程名称'):
                course = i[0]
                df = i[1]
                # print(df)
                generate_jxzlfx(course, df, f_path, obj.template_id, obj.dept_name, obj.major)
            # ftt = shutil.make_archive(f'media/grades_files/{f_path}', 'zip', f'media/grades_files/{f_path}')
            grades_analysis_file = open(f'media/grades_files/{f_path + ".zip"}', 'rb')
            response = FileResponse(grades_analysis_file)
            # response['Content-Type'] = 'application/octet-stream'
            response['Content-Disposition'] = f'attachment;filename="{escape_uri_path(f_path + ".zip")}"'

            return response
        else:
            return super().response_change(request, obj)

    def change_view(self, request, object_id, form_url='', extra_context=None):
        extra_context = extra_context or {}
        extra_context['show_save_and_add_another'] = False
        return super().change_view(request, object_id, form_url, extra_context=extra_context)

    def get_form(self, request, obj=None, **kwargs):
        inst = UserInfo.objects.get(username=request.user)

        form = super().get_form(request, obj=obj, **kwargs)
        form.base_fields['dept_name'].initial = inst.dept_name
        form.base_fields['major'].initial = inst.major
        form.base_fields['template_id'].initial = AnalysisTemplate.objects.filter(instructor=inst)[0]
        return form


    def add_view(self, request, form_url="", extra_context=None):
        extra_context = extra_context or {}
        extra_context['show_save_and_add_another'] = False

        if request.method == 'GET':
            # extra_context = extra_context or {}
            # extra_context['show_save_and_add_another'] = False

            return super(GradesFilesAdmin, self).add_view(request, form_url, extra_context)
        elif request.method == 'POST':
            return self.changeform_view(request, None, form_url, extra_context)


            # 在django中，用session传递redirect的参数
            # request.session["down_file"] = f_path + '.zip'
            # return redirect('/grades/upload_success')

admin.site.register(GradesFiles, GradesFilesAdmin)


# 用于在直方图上标注数字的函数
def autolabel(rects):
    for rect in rects:
        height = rect.get_height()
        plt.text(rect.get_x() + rect.get_width() / 2 - 0.08, 1.03 * height, '%s' % int(height), size=10,
                 family='Times new roman')

# 将系统的 学年学期 转换成 自然年学期
def xuenian_to_zirannian(xn):
    zrn = "年"
    if xn[-1] == '1':
        zrn = xn[:4] + zrn + "下学期"
    elif xn[-1] == '2':
        zrn = xn[5:9] + zrn + "上学期"
    return zrn

# 将传入的文件进行预处理，转换成dataframe
def preprocess(file):
    data = pd.read_excel(file)
    # print(data)
    data.dropna(thresh=3).drop_duplicates()
    grade_data = data.dropna(thresh=3).drop_duplicates()
    grade_data.columns = grade_data.iloc[0]
    grade_data = grade_data[1:]
    # 第1列以后为空的boolean数组  na_bool_index
    na_bool_index = data.dropna(how='all').drop_duplicates().iloc[:, 1].isna()
    # 用 na_bool_index 筛选出第1列为空的第一列值的series
    first_col = data.dropna(how='all').drop_duplicates().iloc[:, 0][na_bool_index]
    # 用以保存第一列值的结果
    first_col_list = []
    for item in first_col:
        tmp_list = re.split(" |,", item)
        if len(tmp_list) > 1:
            for i in tmp_list:
                first_col_list.append(i)
    # print(first_col_list)
    file_info_list = []
    cal_method = ''  # 成绩计算方法
    for item in first_col_list:
        #     print(re.split(":|：", item)[-2].replace("占","占比"))
        k = re.split(":|：", item)[-2].replace("占", "占比")
        v = re.split(":|：", item)[-1]
        grade_data[k] = v
        if '占比' in k:
            cal_method = cal_method + item + '   '
    file_info_list.append(grade_data['课程名称'].iloc[0].replace("/", " "))
    file_info_list.append(grade_data['上课教师'].iloc[0])
    file_info_list.append(grade_data['教学班'].iloc[0])
    file_info_list.append(grade_data['开课学期'].iloc[0])
    file_info_list.append(grade_data['课程类别'].iloc[0])
    file_info_list.append(grade_data['学时'].iloc[0])
    return grade_data, file_info_list, cal_method

# 生成 班级成绩分析单，并得到完整的file_info_list
def generate_cjfxd(grade_data, file_info_list, cal_method, path, analysisTemplate):
    wordname = file_info_list[0] + " " + file_info_list[2] + '班级成绩分析单.docx'  # 成绩分析单名字
    pic_name = file_info_list[0] + " " + file_info_list[2] + '学生成绩正态分布图.png'
    wordfile = Document()
    wordfile.styles['Normal'].font.name = 'Arial'
    wordfile.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    title = wordfile.add_paragraph('湖南商务职业技术学院结课考试班级成绩分析单')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.font.bold = True
    part1 = wordfile.add_paragraph('\n一、不及格、舞弊、缺考名单')
    for run in part1.runs:
        run.font.size = Pt(12)
        run.font.bold = True
    # 打印不及格名单
    if pd.Series(grade_data['总成绩'].astype('float64') < 60).any():  #如果总成绩中有不及格的
        fail_df = grade_data[['学生学号', '学生姓名', '总成绩']][grade_data['总成绩'].astype('float64') < 60]
        t1 = wordfile.add_table(rows=1, cols=3, style='Table Grid')
        t1_cells = t1.rows[0].cells
        t1_cells[0].text = '学生学号'
        t1_cells[1].text = '学生姓名'
        t1_cells[2].text = '总成绩'
        for i in range(fail_df.shape[0]):
            row_cells = t1.add_row().cells
            for j in range(fail_df.shape[1]):
                # row_cells[j].text = str(fail_df.iloc[i][j])
                if isinstance(fail_df.iloc[i][j], float):
                    row_cells[j].text = f'{round(fail_df.iloc[i][j],1)}'
                else:
                    row_cells[j].text = fail_df.iloc[i][j]
        # 设置行高列宽
        for row in t1.rows:
            row.height = Cm(0.7)
        for column in t1.columns:
            for cell in column.cells:
                cell.width = Cm(2)
    else:
        para1 = wordfile.add_paragraph("无")
        para1.paragraph_format.first_line_indent = Pt(22)

    part2 = wordfile.add_paragraph('\n二、成绩计算方法')
    for run in part2.runs:
        run.font.size = Pt(12)
        run.font.bold = True
    para2 = wordfile.add_paragraph(cal_method)
    para2.paragraph_format.first_line_indent = Pt(22)

    part3 = wordfile.add_paragraph('三、学生成绩正态分布图')
    for run in part3.runs:
        run.font.size = Pt(12)
        run.font.bold = True
    wordfile.add_paragraph('1、各档成绩百分比')
    t3 = wordfile.add_table(rows=1, cols=3, style='Table Grid')
    # t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    t3_title = ['项目', '人数', '百分比']
    for t, cell in zip(t3_title, t3.rows[0].cells):
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t3_run = cell.paragraphs[0].add_run(t)
        t3_run.font.bold = True
    project_list = ['应考人数', '缺考人数', '100-90分', '89-80分', '79-70分', '69-60分', '59分及其以下',
                    '平均分']
    person_list = []
    percent_list = ['']
    person_list.append(grade_data.shape[0])
    # 找不到数据，默认第二个元素（缺考人数）为0
    person_list.append(0)
    percent_list.append(0 / grade_data.shape[0])
    grade_range_list = []  # 用以保存分数段的统计，以及逆序的依据，因为cut方法只支持单调递增数组
    # 人数按分数分组统计
    grade_range = pd.cut(grade_data['总成绩'].astype('float64'), [0, 60, 70, 80, 90, 100.1], right=False)
    for i in grade_data['总成绩'].groupby(grade_range):
        #     print(i[0], len(i[1]))
        grade_range_list.append(len(i[1]))
    for i in reversed(grade_range_list):
        person_list.append(i)
        percent_list.append(i / grade_data.shape[0])
    person_list.append(grade_data['总成绩'].astype('float64').mean())
    percent_list.append('')
    # print(grade_range_list[::-1])  #不明白reverse为什么不行，只好用这种方法
    # print(grade_range_list)
    # print(percent_list)
    align_list = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
    for i, j, k in zip(project_list, person_list, percent_list):
        #     必须格式化字符，否则提示错误'float' object is not iterable
        if isinstance(j, int):
            j = str(j)
        else:
            j = f'{j:.2f}'
        if isinstance(k, float):
            k = f'{k * 100:.1f}%'
        row_cells = t3.add_row().cells
        for cell, a, r in zip(row_cells, align_list, [i, j, k]):
            cell.paragraphs[0].paragraph_format.alignment = a
            cell.paragraphs[0].add_run(str(r))
    # 设置t3列宽
    width_list = [4, 2.5, 2]
    for column, w in zip(t3.columns, width_list):
        for cell in column.cells:
            cell.width = Cm(w)
    file_info_list = file_info_list + person_list
    file_info_list.append(pic_name)

    wordfile.add_paragraph('\n2、学生成绩正态分布图')
    mpl.rcParams['font.sans-serif'] = ['SimHei']

    x = np.arange(5)
    y1 = grade_range_list[::-1]
    plt.figure()  # 不加这一句图像会不停覆盖
    bar_draw = plt.bar(x, y1, tick_label=project_list[2:7], width=0.5)
    # print(y1)
    # print(project_list[2:7])
    autolabel(bar_draw)
    plt.xlabel('分段', size=12)
    plt.ylabel('人数', size=12)
    plt.ylim((0, max(y1) * 1.15))
    plt.title(pic_name[:-4])
    plt.savefig(f'media/grades_files/{path}pic/{pic_name}')
    wordfile.add_picture(f'media/grades_files/{path}pic/{pic_name}')

    part4 = wordfile.add_paragraph('四、考试反映的问题及试卷评价')
    for run in part4.runs:
        run.font.size = Pt(12)
        run.font.bold = True
    # para4 = wordfile.add_paragraph('总体来说，学生成绩一般，还可以取得更好的成绩，实际动手能力还有待提高，试卷难度适中。')
    para4 = wordfile.add_paragraph(analysisTemplate.comment)
    para4.paragraph_format.first_line_indent = Pt(22)

    wordfile.save(f'media/grades_files/{path}/{wordname}')
    return file_info_list[0]+file_info_list[2],  file_info_list

# 生成以课程为单位的教学质量分析文档
def generate_jxzlfx(course, df, path, analysisTemplate, dept_name, major):
    jxzlfx_name = course + " " + '教学质量分析.docx'
    jxzlfx = Document()
    jxzlfx.styles['Normal'].font.name = 'Arial'
    jxzlfx.styles['Normal'].font.size = Pt(12)
    jxzlfx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    jxzlfx.add_paragraph()
    para_pic = jxzlfx.add_paragraph()
    para_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pic = para_pic.add_run('')
    run_pic.add_picture('media/hnswxy.png')
    cover_title = jxzlfx.add_paragraph('\n\n\n教　学　质　量　分　析\n\n\n\n')
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cover_title.runs:
        run.font.size = Pt(26)
        run.font.bold = True
    kcmc = jxzlfx.add_paragraph('课程名称：  ')
    kcmc.add_run(course + "  _").underline = True
    kcmc.paragraph_format.first_line_indent = Pt(100)
    for run in kcmc.runs:
        run.font.size = Pt(16)
    sksj = jxzlfx.add_paragraph('授课时间：  ')
    sksj.add_run(xuenian_to_zirannian(df['授课时间'].iloc[0]) + "     _").underline = True
    sksj.paragraph_format.first_line_indent = Pt(100)
    for run in sksj.runs:
        run.font.size = Pt(16)
    rkjs = jxzlfx.add_paragraph('任课教师： ')
    rkjs.add_run(df['上课教师'].iloc[0] + '                      _').underline = True
    rkjs.paragraph_format.first_line_indent = Pt(100)
    for run in rkjs.runs:
        run.font.size = Pt(16)
    yb = jxzlfx.add_paragraph('院       部：  ')
    yb.add_run(dept_name + '       _').underline = True
    yb.paragraph_format.first_line_indent = Pt(100)
    for run in yb.runs:
        run.font.size = Pt(16)
    jys = jxzlfx.add_paragraph('教  研  室：  ')
    jys.add_run(major + '       _').underline = True
    jys.paragraph_format.first_line_indent = Pt(100)
    for run in jys.runs:
        run.font.size = Pt(16)
    # 规范页
    jxzlfx.add_page_break()
    guifan_title = jxzlfx.add_paragraph('\n教学质量分析规范\n\n')
    guifan_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in guifan_title.runs:
        run.font.size = Pt(16)
    gf_p1 = jxzlfx.add_paragraph(
        '一、教学质量分析是教师自我总结教学情况，学校有关部门了解教学情况的基本依据，每位教师务必严肃、认真、规范地搞好质量分析。')
    gf_p2 = jxzlfx.add_paragraph('二、教学质量分析至少每学期进行一次，由任课教师按课程撰写。')
    gf_p3 = jxzlfx.add_paragraph('三、教学质量分析的基本内容')
    gf_p4 = jxzlfx.add_paragraph('（一）班级教学情况分析')
    gf_p5 = jxzlfx.add_paragraph('1、学生基本情况')
    gf_p6 = jxzlfx.add_paragraph('（包括各班人数，上课、自习，回答问题，作业、实验、实训，考试等环节的表现）')
    gf_p7 = jxzlfx.add_paragraph('2、教师教学情况')
    gf_p8 = jxzlfx.add_paragraph('（包括因人、因材施教，完成教学任务以及教学态度、能力、效果等方面）')
    gf_p9 = jxzlfx.add_paragraph('（二）取得成功的经验及失误的教训分析')
    gf_p10 = jxzlfx.add_paragraph('（三）合理化建议\n\n\n')
    for p in [gf_p1, gf_p2, gf_p3, gf_p4, gf_p5, gf_p6, gf_p7, gf_p8, gf_p9, gf_p10]:
        p.paragraph_format.first_line_indent = Pt(24)
        p.line_spacingpacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(22)
        p.paragraph_format.space_after = Pt(0)
    gf_p11 = jxzlfx.add_paragraph('教务处')
    gf_p11.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # 正文页
    jxzlfx.add_page_break()
    zw_title = jxzlfx.add_paragraph('教学质量分析\n')
    zw_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in zw_title.runs:
        run.font.size = Pt(24)
    zw1 = jxzlfx.add_paragraph('撰写人： ')
    zw1.add_run(df['上课教师'].iloc[0])
    zw1.add_run('        撰写时间：')
    zw1.add_run(pd.Timestamp.today().strftime('%Y年%m月%d日'))
    zw1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in zw1.runs:
        run.font.size = Pt(16)
    t_zw2 = jxzlfx.add_table(rows=1, cols=1, style='Table Grid')
    t_zw2.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 通过默认实训的课时计算方法初始化
    js = int(df['学时'].iloc[0])//10
    sy = int(df['学时'].iloc[0]) - js
    cy = 0
    jd = 0  #机动课时
    if '实训' not in course:
        if int(df['学时'].iloc[0]) / 10 >= 8:
            jd = 8
        elif int(df['学时'].iloc[0]) / 10 < 4:
            jd = 2
        else:
            jd = 4
        js = sy = int((int(df['学时'].iloc[0]) - jd) / 2)
    t_zw2_p1 = t_zw2.cell(0, 0).add_paragraph('一、课程总课时：')
    t_zw2_p1.add_run(df['学时'].iloc[0]).underline = True
    t_zw2_p1.add_run('    本学期授课时数：')
    t_zw2_p1.add_run(df['学时'].iloc[0]).underline = True
    t_zw2_p1.add_run('    其中讲授')
    t_zw2_p1.add_run(str(js)).underline = True
    t_zw2_p1.add_run('课时，实验（训）')
    t_zw2_p1.add_run(str(sy)).underline = True
    t_zw2_p1.add_run('课时，测验')
    t_zw2_p1.add_run(str(cy)).underline = True
    t_zw2_p1.add_run('课时，机动')
    t_zw2_p1.add_run(str(jd)).underline = True
    t_zw2_p1.add_run('课时')
    t_zw2_p2 = t_zw2.cell(0, 0).add_paragraph('二、教材版本：高职高专教材')
    t_zw2_p3 = t_zw2.cell(0, 0).add_paragraph('三、学期总评成绩：')
    t_zw2_p4 = t_zw2.cell(0, 0).add_paragraph('同课程学生人数：')
    t_zw2_p4.add_run(str(sum(df['应考人数'].astype(int)))).underline = True
    t_zw2_p4.add_run('      平均分：')
    average = np.mean(df['平均分'].astype(float))
    t_zw2_p4.add_run(f'{average:.2f}').underline = True
    t_zw2_p5 = t_zw2.cell(0, 0).add_paragraph('90分以上：')
    t_zw2_p5.add_run(str(sum(df['100-90分'].astype(int)))).underline = True
    t_zw2_p5.add_run('人    80分以上：')
    t_zw2_p5.add_run(str(sum(df['89-80分'].astype(int)))).underline = True
    t_zw2_p5.add_run('人    70分以上：')
    t_zw2_p5.add_run(str(sum(df['79-70分'].astype(int)))).underline = True
    t_zw2_p5.add_run('人')
    t_zw2_p6 = t_zw2.cell(0, 0).add_paragraph('60分以上：')
    t_zw2_p6.add_run(str(sum(df['69-60分'].astype(int)))).underline = True
    t_zw2_p6.add_run('人    不及格：')
    t_zw2_p6.add_run(str(sum(df['59分及其以下'].astype(int)))).underline = True
    t_zw2_p6.add_run('人\n')
    for p in [t_zw2_p1, t_zw2_p2, t_zw2_p3, t_zw2_p4, t_zw2_p5, t_zw2_p6]:
        p.paragraph_format.first_line_indent = Pt(24)
        p.line_spacingpacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(22)
        p.paragraph_format.space_after = Pt(0)

    zw3 = jxzlfx.add_paragraph('\n成绩分析柱状图：')
    for i in df['图片名']:
        jxzlfx.add_picture(f'media/grades_files/{path}pic/{i}')
    fx1 = jxzlfx.add_paragraph('\n（一）班级教学情况分析')
    fx2 = jxzlfx.add_paragraph('1、学生基本情况')
    fx3 = jxzlfx.add_paragraph('')
    for i, j in zip(df['班级名称'], df['应考人数']):
        fx3.add_run(i)
        fx3.add_run('班一共')
        fx3.add_run(str(j))
        fx3.add_run('人，')
    fx3.add_run(analysisTemplate.basic_condition)
    # fx3.add_run('上课期间大部分学生上课都比较积极和认真，')
    # if '实训' in course:
    #     fx3.add_run('而由于实训以学生动手操作为主，学生也因此体现出更高的积极性，在做实训项目时，学生遇到问题都会及时向老师询问。上机以学生提交的实训项目材料为考察依据，所有同学按时提交了考试相关材料，并留有存档。')
    # else:
    #     fx3.add_run('在做例题和练习时，学生遇到问题都会及时向老师询问。考查采用学习通题库随机组卷的方式。')
    fx4 = jxzlfx.add_paragraph('2、教师教学情况')
    fx5 = jxzlfx.add_paragraph(analysisTemplate.teaching_condition)
    # fx5 = jxzlfx.add_paragraph(
    #     '第一，积极向教学丰富的教师请教教学经验；第二，上课前认真备课，并且坚持课前演练；第三，针对学生素质不等，所以授课时要考虑整体情况，因材施教，既不能太快也不能太慢，更加不能太枯燥。')
    fx6 = jxzlfx.add_paragraph('（二）取得成功的经验及失误的教训分析')
    fx7 = jxzlfx.add_paragraph(analysisTemplate.experience)
    # fx7 = jxzlfx.add_paragraph(
    #     '为了保证学生能够学会项目相关的知识，一定要在每讲完一个知识点后给学生理解和练习的时间，教师要走下讲堂，认真发现学生的问题，对于那些不会动手的学生要手把手教，把他们及时入门。')
    fx8 = jxzlfx.add_paragraph('（三）今后改进教学的措施')
    fx9 = jxzlfx.add_paragraph(analysisTemplate.advice.replace('\r\n','\n\t'))
    # fx9 = jxzlfx.add_paragraph('1、学生的思想教育很重要，技术再厉害，思想也要先行。')
    # fx10 = jxzlfx.add_paragraph(
    #     '2、通过从提交的结果发现，在今后的教学过程中，需要更加注意细节；另外就是要注重逻辑思维的养成以及分析和解决问题的能力培养。')
    # fx11 = jxzlfx.add_paragraph('3、培养学生的实际动手能力很重要，要充分调动学生学习的积极性。')
    for p in [fx1, fx2, fx3, fx4, fx5, fx6, fx7, fx8, fx9]:
    # for p in [fx1, fx2, fx3, fx4, fx5, fx6, fx7, fx8, fx9, fx10, fx11]:
        p.paragraph_format.first_line_indent = Pt(24)
        p.line_spacingpacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(22)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.underline = True
    jxzlfx.save(f'media/grades_files/{path}/{jxzlfx_name}')