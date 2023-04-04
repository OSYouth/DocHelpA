import os, shutil

from django.shortcuts import render
from django.http import HttpResponse, FileResponse, StreamingHttpResponse
from .models import *
from django.contrib.auth.decorators import login_required
from user.models import UserInfo
from docx import Document
from docx.shared import Pt,Cm    #字号
from docx.oxml.ns import qn    #字体
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING   #对齐
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT    #对齐
import time
import pandas as pd
import numpy as np
from .add_float_picture import add_float_picture

# Create your views here.
# 生成要下载的压缩包
def download(request):
    # 根据指导老师、时间，按学号创建文件夹，并对每个学生生成4个文件
    # 根据指导老师、时间，按文件名创建文件夹，并生成指导老师名下所有的学生的文件
    # name = generate_zdjlb('201904040111')
    # name = generate_dbjlb('201904040111', '00530')
    # name = generate_pyb('201904040111', '00530')
    name = generate_rws('201904040111', '00530')
    return render(request, 'upload_success.html', locals())

@login_required
def download_files(request, by):
    if not os.path.exists(f'media/download_files'):
        os.mkdir(f'media/download_files')
    # 请求账户所有文件保存的路径
    f_path = f'media/download_files/{request.user.username+request.user.last_name+request.user.first_name}'
    if not os.path.exists(f_path):
        os.mkdir(f_path)
    instructor_number = request.user.username
    # 已毕业的学生也要过滤掉，根据User_Info.title != 'ungraduate'进行判断
    student_charge_set = GraduateProjectInfo.objects.filter(instructor=request.user.last_name+request.user.first_name).filter(stu__title="ungraduate")
    if by == 'student_number':
        # 存储此次所有文件 的文件夹的路径，名字不同以秒间隔开来，所以没有做路径是否存在的判断
        filename = f'{request.user.last_name + request.user.first_name}（按学号创建）{pd.Timestamp.today().strftime("%Y%m%d%H%M%S")}'
        downfile_path = f_path +"/" +  filename
        os.mkdir(downfile_path)
        for s in student_charge_set:
            stu_file_path = downfile_path + f'/{s.stu.username+s.stu.last_name+s.stu.first_name}'
            os.mkdir(stu_file_path)
            generate_rws(  s.stu.username, instructor_number, stu_file_path)
            generate_zdjlb(s.stu.username, instructor_number, stu_file_path)
            generate_dbjlb(s.stu.username, instructor_number, stu_file_path)
            generate_pyb(  s.stu.username, instructor_number, stu_file_path)
        shutil.make_archive(f'{downfile_path }', 'zip', f'{downfile_path }')
        response = FileResponse(file_iterator(f'{downfile_path}.zip'))
        response['Content-Type'] = 'application/octet-stream'
        response['Content-Disposition'] = f'attachment;filename={filename.encode("utf-8").decode("ISO-8859-1")}.zip'
        return response
    elif by == 'file_type':
        # 存储此次所有文件 的文件夹的路径，名字不同以秒间隔开来，所以没有做路径是否存在的判断
        filename = f'{request.user.last_name + request.user.first_name}（按文件创建）{pd.Timestamp.today().strftime("%Y%m%d%H%M%S")}'
        downfile_path = f_path + "/" + filename
        os.mkdir(downfile_path)

        for s in student_charge_set:
            rws_path = downfile_path + f'/0.毕业设计任务书'
            if not os.path.exists(rws_path):
                os.mkdir(rws_path)
            generate_rws(  s.stu.username, instructor_number, rws_path)
            zdjlb_path = downfile_path + f'/1.毕业设计指导记录表'
            if not os.path.exists(zdjlb_path):
                os.mkdir(zdjlb_path)
            generate_zdjlb(s.stu.username, instructor_number, zdjlb_path)
            pyb_path = downfile_path + f'/2.毕业设计评阅表'
            if not os.path.exists(pyb_path):
                os.mkdir(pyb_path)
            generate_pyb(s.stu.username, instructor_number, pyb_path)
            dbjlb_path = downfile_path + f'/3.毕业设计答辩记录表'
            if not os.path.exists(dbjlb_path):
                os.mkdir(dbjlb_path)
            generate_dbjlb(s.stu.username, instructor_number, dbjlb_path)

        shutil.make_archive(f'{downfile_path }', 'zip', f'{downfile_path }')
        response = FileResponse(file_iterator(f'{downfile_path}.zip'))
        response['Content-Type'] = 'application/octet-stream'
        response['Content-Disposition'] = f'attachment;filename={filename.encode("utf-8").decode("ISO-8859-1")}.zip'
        return response

def file_iterator(file_path, chunk_size=512):
    """
    文件生成器,防止文件过大，导致内存溢出
    :param file_path: 文件绝对路径
    :param chunk_size: 块大小
    :return: 生成器
    """
    with open(file_path, mode='rb') as f:
        while True:
            c = f.read(chunk_size)
            if c:
                yield c
            else:
                break

def download_test(request, name):
    file = open(f'media/download_files/{name}', 'rb')
    response = FileResponse(file)
    response['Content-Type'] = 'application/octet-stream'
    response['Content-Disposition'] = f'attachment;filename={name}'
    return response

# 生成毕业设计任务书
def generate_rws(stu_no, inst_no, path):
    stu = UserInfo.objects.get(username=stu_no)
    inst = UserInfo.objects.get(username=inst_no)
    stu_project = GraduateProjectInfo.objects.get(stu=stu)
    stu_defence = DefenceInfo.objects.get(stu=stu)
    temp_set = AssignmentTemplate.objects.filter(instructor=inst)
    temp = temp_set[0]  # 默认使用教师名下的第一个模板
    temp_set_id_list = list(AssignmentTemplate.objects.filter(instructor=inst).values_list('id', flat=True))
    # 判断学生答辩信息中模板id是否在 获取所有模板的ID列表 中
    if stu_defence.assignment_template_id in list(map(str,temp_set_id_list)):
        temp = temp_set.get(id=stu_defence.assignment_template_id)
    word_name = stu_no + "+" + stu_project.class_name + "+" + stu.last_name + stu.first_name + "+毕业设计任务书.docx"
    word = Document()
    # 先对于文档英文为Arial，中文为宋体所有字体设置为小四，所有字体设置为小四，
    word.styles['Normal'].font.name = 'Arial'
    word.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    word.styles['Normal'].font.size = Pt(12)
    word.sections[0].left_margin = Cm(2.1)
    word.sections[0].right_margin = Cm(2.1)
    title = word.add_paragraph('湖南商务职业技术学院毕业设计任务书')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.font.bold = True
    table = word.add_table(rows=9, cols=6, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头前三行对齐格式统一
    align_cccccl = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    align_clcl = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    row1 = []    #第1行的元素值用列表存储
    row1.append('学生姓名')
    row1.append(stu.last_name + stu.first_name)
    row1.append('二级学院')
    row1.append(stu.dept_name)
    row1.append('班级名称')
    row1.append(stu_project.class_name)
    for cell, align, ele in zip(table.rows[0].cells, align_cccccl, row1):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.alignment = align
        cell.paragraphs[0].add_run(ele)
    row2 = []  # 第2行的元素值用列表存储
    row2.append('专业名称')
    row2.append(stu.major)
    row2.append('学      号')
    row2.append(stu_no)
    row2.append('指导教师')
    row2.append(stu_project.instructor)
    for cell, align, ele in zip(table.rows[1].cells, align_cccccl, row2):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.alignment = align
        cell.paragraphs[0].add_run(ele)
    row3 = []  # 第3行的元素值用列表存储
    row3.append("选题名称")
    row3.append('')
    row3.append('')
    row3.append(stu_project.topic)
    row3.append("设计类型")
    row3.append(stu_defence.design_type)
    table.rows[2].cells[1].merge(table.rows[2].cells[3])
    for cell, align, ele in zip(table.rows[2].cells, (align_clcl*2)[:6], row3):
        cell.paragraphs[0].paragraph_format.alignment = align
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].add_run(ele)
    # 从第四行开始，全部都只有一个单元格
    for i in range(3,9):
        table.rows[i].cells[0].merge(table.rows[i].cells[5])
    table.rows[3].cells[0].paragraphs[0].add_run('一、毕业设计目标')
    keyword_length = temp.goal.count("*")
    keyword_list = [stu_project.key_word1, stu_project.key_word2, stu_project.key_word3, stu_project.key_word4, stu_project.key_word5]
    design_object = temp.goal.replace("*","%s")%tuple(keyword_list[:keyword_length])
    table.rows[3].cells[0].add_paragraph().add_run('\t'+design_object)

    table.rows[4].cells[0].paragraphs[0].add_run('二、毕业设计任务及要求')
    table.rows[4].cells[0].add_paragraph().add_run('\t'+temp.task_require.replace('\r\n','\n\t'))

    table.rows[5].cells[0].paragraphs[0].add_run('三、毕业设计实施步骤和方法')
    table.rows[5].cells[0].add_paragraph().add_run('\t'+temp.step_way.replace('\r\n','\n\t'))

    table.rows[6].cells[0].paragraphs[0].add_run('四、毕业设计进程安排')
    table.rows[6].cells[0].add_paragraph().add_run('\t'+temp.schdeule.replace(" ","").replace('\n','\n\t'))

    table.rows[7].cells[0].paragraphs[0].add_run('五、设计思路')
    table.rows[7].cells[0].add_paragraph().add_run('\t'+temp.thought.replace('\r\n','\n\t'))

    table.rows[8].cells[0].paragraphs[0].add_run('六、成果表现形式')
    table.rows[8].cells[0].add_paragraph().add_run('\t'+temp.result.replace('\r\n','\n\t'))

    # 保证前面创建的table格式里面的行高至少为1cm
    for i in range(9):
        table.rows[i].height = Cm(1)
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = Pt(20)
                p.paragraph_format.space_after = Pt(0)

    table2 = word.add_table(rows=2, cols=14, style='Table Grid')
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.rows[0].cells[1].merge(table2.rows[0].cells[13])
    table2.rows[0].cells[0].paragraphs[0].add_run('指导教师意见')
    table2.rows[0].cells[1].add_paragraph().add_run(temp.comment)
    p10 = table2.rows[0].cells[1].add_paragraph('签字：')
    try:
        p10.add_run().add_picture(f'media/{inst.sign}', width=Cm(1.8))
    except Exception:
        p10.add_run().add_picture(f'static/缺失.jpeg', width=Cm(1.8))
        p10.add_run(inst.last_name + inst.first_name + "信息不完整，请完善信息")
    p10.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table2.rows[0].cells[1].add_paragraph(temp.rws_date).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    table2.rows[1].cells[1].merge(table2.rows[1].cells[6])
    table2.rows[1].cells[8].merge(table2.rows[1].cells[13])
    table2.rows[1].cells[0].paragraphs[0].add_run('教研室主任审批意见')
    table2.rows[1].cells[1].add_paragraph('      同意实施')
    table2.rows[1].cells[1].add_paragraph()
    p11_1 = table2.rows[1].cells[1].add_paragraph('签字：')
    try:
        director = UserInfo.objects.get(last_name=stu_project.director[0], first_name=stu_project.director[1:])
        p11_1.add_run().add_picture(f'media/{director.sign}', width=Cm(1.8))
    except Exception:
        p11_1.add_run().add_picture(f'static/缺失.jpeg', width=Cm(1.8))
        p11_1.add_run(director.last_name + director.first_name + "信息不完整，请联系其完善信息")
    p11_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table2.rows[1].cells[1].add_paragraph(temp.rws_date).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table2.rows[1].cells[7].paragraphs[0].add_run('学院审批意见')
    table2.rows[1].cells[7].paragraphs[0].width = Cm(2.1)
    table2.rows[1].cells[8].add_paragraph('      同意实施')
    table2.rows[1].cells[8].add_paragraph()
    p11_2 = table2.rows[1].cells[8].add_paragraph('签字：')
    p11_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    try:
        dean = UserInfo.objects.get(last_name=stu_project.dean[0], first_name=stu_project.dean[1:])
        p11_2.add_run().add_picture(f'media/{dean.sign}', width=Cm(1.8))
    except Exception:
        p11_2.add_run().add_picture(f'static/缺失.jpeg', width=Cm(1.8))
        p11_2.add_run(dean.last_name + dean.first_name + "信息不完整，请联系其完善信息")
    add_float_picture(p11_2, f'media/{temp.cachet}', width=Cm(4.2), pos_x=Cm(1.5), pos_y=Cm(-0.3))
    table2.rows[1].cells[8].add_paragraph(temp.rws_date).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for i in range(2):
        table2.rows[i].height = Cm(1)

    # p11_2.add_run().add_picture(f'media/{temp.cachet}', width=Cm(4))
    # p = word.add_paragraph()
    # add_float_picture(p, f'media/{temp.cachet}', width=Cm(3.6), pos_x=Cm(20), pos_y=Cm(30))

    bz1 = word.add_paragraph('注：1. 设计类型包括产品设计、流程设计或方案设计。')
    bz2 = word.add_paragraph('       2. 本表一式两份，一份二级学院留存，一份发给学生。')
    bz3 = word.add_paragraph('       3. 所有签字处必须手写，其他内容可打印，未签字者，不予开题。')
    for p in [bz1, bz2, bz3]:
        p.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.size = Pt(10.5)

    word.save(f'{path}/{word_name}')
    return word_name

# 生成毕业设计指导记录表
def generate_zdjlb(stu_no, inst_no, path):
    inst = UserInfo.objects.get(username=inst_no)
    stu = UserInfo.objects.get(username=stu_no)
    stu_project = GraduateProjectInfo.objects.get(stu=stu)
    stu_defence = DefenceInfo.objects.get(stu=stu)
    grtemp = GuideRecordTemplate.objects.get(instructor=inst)
    word_name =stu_no + "+" + stu_project.class_name + "+" + stu.last_name + stu.first_name + "+毕业设计指导记录表.docx"
    word = Document()
    # 先对于文档英文为Arial，中文为宋体，所有字体设置为小四
    word.styles['Normal'].font.name = 'Arial'
    word.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    word.styles['Normal'].font.size = Pt(12)
    title = word.add_paragraph('湖南商务职业技术学院毕业设计指导记录表')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.font.bold = True
    table = word.add_table(rows=3, cols=6, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头前三行对齐格式统一
    align_cccccl = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    align_clcl = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    row1 = []    #第1行的元素值用列表存储
    row1.append('学生姓名')
    row1.append(stu.last_name + stu.first_name)
    row1.append('二级学院')
    row1.append(stu.dept_name)
    row1.append('班级名称')
    row1.append(stu_project.class_name)
    for cell, align, ele in zip(table.rows[0].cells, align_cccccl, row1):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.alignment = align
        cell.paragraphs[0].add_run(ele)
    row2 = []  # 第2行的元素值用列表存储
    row2.append('专业名称')
    row2.append(stu.major)
    row2.append('学      号')
    row2.append(stu_no)
    row2.append('指导教师')
    row2.append(stu_project.instructor)
    for cell, align, ele in zip(table.rows[1].cells, align_cccccl, row2):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.alignment = align
        cell.paragraphs[0].add_run(ele)
    row3 = []  # 第3行的元素值用列表存储
    row3.append("选题名称")
    row3.append('')
    row3.append('')
    row3.append(stu_project.topic)
    row3.append("设计类型")
    row3.append(stu_defence.design_type)
    table.rows[2].cells[1].merge(table.rows[2].cells[3])
    for cell, align, ele in zip(table.rows[2].cells, (align_clcl*2)[:6], row3):
        cell.paragraphs[0].paragraph_format.alignment = align
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].add_run(ele)
    for i in range(3):
        table.rows[i].height = Cm(1)

    table2 = word.add_table(rows=8, cols=5, style='Table Grid')
    # 统一将第2个表格的第4列和第5列合并
    for i in range(8):
        table2.rows[i].cells[3].merge(table2.rows[i].cells[4])
        table2.rows[i].height = Cm(1)
    # 合并作为毕业设计过程的单元格
    table2.rows[3].cells[0].merge(table2.rows[6].cells[0])
    row4 = []  # 第4行的元素值用列表存储    默认第5个元素和第4个一样，不做处理
    row4.append("指导内容")
    row4.append('指导时间')
    row4.append('指导地点')
    row4.append("指导过程记录")
    for cell, align, ele in zip(table2.rows[0].cells[:4], align_cccccl[:4], row4):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.alignment = align
        cell.paragraphs[0].add_run(ele)
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(14)
            run.font.bold = True

# 从这一行开始后面的6行可以用for循环结合数据库进行遍历
    # 在进入程序开始需要引入老师的信息以调用指导记录表的模板信息
    # 还需要注意日期的年份要在程序里面重新获取，以免模板在教师修改后，年份写死了
    row5to10 = [ grtemp.guide_cont1, grtemp.guide_date1, grtemp.guide_loca1, grtemp.guide_proc1, grtemp.guide_cont2, grtemp.guide_date2, grtemp.guide_loca2, grtemp.guide_proc2, grtemp.guide_cont3, grtemp.guide_date3, grtemp.guide_loca3, grtemp.guide_proc3, grtemp.guide_cont3_2, grtemp.guide_date3_2, grtemp.guide_loca3_2, grtemp.guide_proc3_2, grtemp.guide_cont3_3, grtemp.guide_date3_3, grtemp.guide_loca3_3, grtemp.guide_proc3_3, grtemp.guide_cont3_4, grtemp.guide_date3_4, grtemp.guide_loca3_4, grtemp.guide_proc3_4]
    for i in range(6):
        for cell, align, ele in zip(table2.rows[i+1].cells[:4], align_clcl, np.array(row5to10).reshape(6,4)[i]):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].paragraph_format.alignment = align
            try:
                d = time.strptime(ele, '%Y年%m月%d日')
                # 月份大过8则年份肯定是上一年
                if d.tm_mon > 8:
                    cell.paragraphs[0].add_run(f'{time.localtime().tm_year-1}年{d.tm_mon}月{d.tm_mday}日')
                else:
                    cell.paragraphs[0].add_run(f'{time.localtime().tm_year}年{d.tm_mon}月{d.tm_mday}日')
            except Exception:
                cell.paragraphs[0].add_run(ele)

    row11 = []  # 第11行的元素值用列表存储    第1列置空，默认第5个元素和第4个一样，不做处理
    row11.append('毕业设计答辩')
    row11.append(stu_defence.defence_date)
    row11.append('线上')
    row11.append("答辩")
    align_clcl = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    for cell, align, ele in zip(table2.rows[7].cells, align_clcl, row11):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.alignment = align
        cell.paragraphs[0].add_run(ele)

    bz1 = word.add_paragraph('注：1. 设计类型包括产品设计、流程设计或方案设计。')
    bz2 = word.add_paragraph('       2. 毕业设计过程部分，指导教师可根据实际情况自行增加或删除行。')
    for p in [bz1, bz2]:
        p.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.size = Pt(10.5)

    word.save(f'{path}/{word_name}')
    return word_name


# 生成毕业设计评阅表
def generate_pyb(stu_no, inst_no, path):
    stu = UserInfo.objects.get(username=stu_no)
    inst = UserInfo.objects.get(username=inst_no)
    stu_project = GraduateProjectInfo.objects.get(stu=stu)
    stu_defence = DefenceInfo.objects.get(stu=stu)
    word_name = stu_no + "+" + stu_project.class_name + "+" + stu.last_name + stu.first_name + "+毕业设计评阅表.docx"
    word = Document()
    # 先对于文档英文为Arial，中文为宋体所有字体设置为小四，所有字体设置为小四，
    word.styles['Normal'].font.name = 'Arial'
    word.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    word.styles['Normal'].font.size = Pt(12)
    word.sections[0].left_margin = Cm(2.1)
    word.sections[0].right_margin = Cm(2.1)
    title = word.add_paragraph('湖南商务职业技术学院毕业设计评阅表')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.font.bold = True
    table = word.add_table(rows=7, cols=6, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头前三行对齐格式统一
    align_cccccl = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    align_clcl = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    row1 = []    #第1行的元素值用列表存储
    row1.append('学生姓名')
    row1.append(stu.last_name + stu.first_name)
    row1.append('二级学院')
    row1.append(stu.dept_name)
    row1.append('班级名称')
    row1.append(stu_project.class_name)
    for cell, align, ele in zip(table.rows[0].cells, align_cccccl, row1):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.alignment = align
        cell.paragraphs[0].add_run(ele)
    row2 = []  # 第2行的元素值用列表存储
    row2.append('专业名称')
    row2.append(stu.major)
    row2.append('学      号')
    row2.append(stu_no)
    row2.append('指导教师')
    row2.append(stu_project.instructor)
    for cell, align, ele in zip(table.rows[1].cells, align_cccccl, row2):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.alignment = align
        cell.paragraphs[0].add_run(ele)
    row3 = []  # 第3行的元素值用列表存储
    row3.append("选题名称")
    row3.append('')
    row3.append('')
    row3.append(stu_project.topic)
    row3.append("设计类型")
    row3.append(stu_defence.design_type)
    table.rows[2].cells[1].merge(table.rows[2].cells[3])
    for cell, align, ele in zip(table.rows[2].cells, (align_clcl*2)[:6], row3):
        cell.paragraphs[0].paragraph_format.alignment = align
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].add_run(ele)
    for i in range(3):
        table.rows[i].height = Cm(1)

    table.rows[3].cells[0].merge(table.rows[3].cells[5])
    table.rows[3].cells[0].paragraphs[0].add_run('评价标准').bold =True
    table.rows[3].cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    row4_1 = table.rows[3].cells[0].add_paragraph()
    row4_1.add_run('1. 选题：').bold = True
    row4_1.add_run('重点评价毕业设计选题的专业性、实践性和工作量。')
    row4_2 = table.rows[3].cells[0].add_paragraph('')
    row4_2.add_run('2. 设计实施：').bold = True
    row4_2.add_run('评价毕业设计项目实施中技术路线的可行性、设计过程的完整性和设计依据的可靠性；按期圆满完成规定的任务，工作量饱满，难度较大；工作努力，遵守纪律；工作作风严谨务实。')
    row4_3 = table.rows[3].cells[0].add_paragraph('')
    row4_3.add_run('3. 分析与解决问题的能力：').bold = True
    row4_3.add_run('能运用所学知识和技能去发现与解决实际问题；能对设计进行理论分析，得出有价值的结论。')
    row4_4 = table.rows[3].cells[0].add_paragraph('')
    row4_4.add_run('4. 成果质量：').bold = True
    row4_4.add_run('以学生毕业设计形成的最终技术文件为主要考察对象，重点评价设计技术文件的规范性、技术方案的科学性和技术设计的创新性。')
    row4_5 = table.rows[3].cells[0].add_paragraph('')
    row4_5.add_run('5. 答辩情况：').bold = True
    row4_5.add_run('阐述课题的设计思路、主要依据、结论、体会和改进意见；回答问题的准确性、敏锐性、全面性、语言表达能力、逻辑条理性。')
    # 对于前5段统一将字体设置为10.5磅
    for i in range(1,6):
        for run in table.rows[3].cells[0].paragraphs[i].runs:
            run.font.size = Pt(10.5)
    row4_6 = table.rows[3].cells[0].add_paragraph('')
    tmp4 = row4_6.add_run('说明：')
    tmp4.bold = True
    tmp4.font.size = Pt(10.5)
    row4_6.add_run('1）指导老师根据标准1—4给出成果成绩').font.size = Pt(10.5)
    row4_6.add_run('（选题10分，设计实施30分，分析与解决问题的能力10分，成果质量50分）')
    row4_6.add_run('，答辩组根据标准1—5给出答辩成绩').font.size = Pt(10.5)
    row4_6.add_run('（选题10分，设计实施20分，分析与解决问题的能力10分，成果质量30分，答辩情况30分）')
    row4_6.add_run('，成果成绩和答辩成绩均采用百分制。').font.size = Pt(10.5)
    row4_7 = table.rows[3].cells[0].add_paragraph('')
    row4_7.add_run('         2）毕业设计最终成绩：成果成绩占40%，答辩成绩占60%，采用五级制').font.size = Pt(10.5)
    row4_7.add_run('（优秀：90以上，良：80-89，中：70-79，及格：60-69，不及格：60分以下）')
    row4_7.add_run('进行评定。').font.size = Pt(10.5)
    for p in table.rows[3].cells[0].paragraphs:
        p.paragraph_format.first_line_indent = Pt(21)
        p.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        # p.paragraph_format.line_spacing = Pt(12)
        p.paragraph_format.space_after = Pt(0)

    table.rows[4].cells[0].merge(table.rows[4].cells[5])
    table.rows[4].cells[0].paragraphs[0].add_run('指导老师评价意见，并对是否同意参加答辩作出明确说明\n').bold = True
    row5_2 = table.rows[4].cells[0].add_paragraph('')
    row5_2.add_run('文档撰写规范，方案设计完善，格式符合要求，同意答辩。\n')
    row5_3 = table.rows[4].cells[0].add_paragraph('')
    tmp5 = row5_3.add_run(' 指导老师签名：')
    try:
        tmp5.add_picture(f'media/{inst.sign}', width=Cm(1.8))
    except Exception:
        tmp5.add_picture(f'static/缺失.jpeg', width=Cm(1.8))
        tmp5.add_text(inst.last_name + inst.first_name + "信息不完整，请完善信息")
    # row5_3.paragraph_format.first_line_indent = Pt(150)
    row5_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tmp_day = pd.to_datetime(stu_defence.defence_date, format='%Y年%m月%d日')-pd.tseries.offsets.DateOffset(days=1)
    row5_4 = table.rows[4].cells[0].add_paragraph(tmp_day.strftime('%Y年%m月%d日'))
    row5_4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.rows[4].cells[0].add_paragraph()

    table.rows[5].cells[0].merge(table.rows[5].cells[5])
    table.rows[5].cells[0].paragraphs[0].add_run('答辩组评价意见\n').bold = True
    row6_2 = table.rows[5].cells[0].add_paragraph('')
    row6_2.add_run('问题回答准确，条理清晰，答辩通过。\n')
    row6_3 = table.rows[5].cells[0].add_paragraph('')
    tmp6 = row6_3.add_run('答辩组老师（三人以上）签名：')
    inst_tmp = [stu_defence.def_inst1, stu_defence.def_inst2, stu_defence.def_inst3, stu_defence.def_inst4, stu_defence.def_inst5]
    for i in inst_tmp:
        # 如果def_inst * 第一个元素（即姓氏）都无法在UserInfo里面找到，则不需要处理
        if len(UserInfo.objects.filter(last_name=i[0])) > 0:
            try:
                tmp_inst = UserInfo.objects.get(last_name=i[0], first_name=i[1:])
                tmp6.add_picture(f'media/{tmp_inst.sign}', width=Cm(1.8))
            except Exception:
                tmp6.add_picture(f'static/缺失.jpeg', width=Cm(1.8))
                tmp6.add_text(tmp_inst.last_name + tmp_inst.first_name + "信息不完整，请联系其完善信息")
    row6_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row6_4 = table.rows[5].cells[0].add_paragraph(stu_defence.defence_date)
    row6_4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.rows[5].cells[0].add_paragraph()

    align_cccccc = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                    WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]

    row7 = []  # 第7行的元素值用列表存储
    row7.append('成果成绩')
    row7.append(stu_defence.ach_grade)
    row7.append('答辩成绩')
    row7.append(stu_defence.defence_grade)
    row7.append('最终成绩')
    row7.append(stu_defence.final_grade)
    for cell, align, ele in zip(table.rows[6].cells, align_cccccc, row7):
        cell.paragraphs[0].add_run(ele)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.alignment = align
    table.rows[6].height = Cm(1)

    bz = word.add_paragraph('注：设计类型指产品设计、流程设计或方案设计。')
    bz.paragraph_format.line_spacing = Pt(15)
    bz.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    for run in bz.runs:
        run.font.size = Pt(10.5)

    word.save(f'{path}/{word_name}')
    return word_name

# 生成毕业设计答辩记录表
def generate_dbjlb(stu_no, inst_no, path):
    stu = UserInfo.objects.get(username=stu_no)
    inst = UserInfo.objects.get(username=inst_no)
    stu_project = GraduateProjectInfo.objects.get(stu=stu)
    stu_defence = DefenceInfo.objects.get(stu=stu)
    word_name = stu_no + "+" + stu_project.class_name + "+" + stu.last_name + stu.first_name + "+毕业设计答辩记录表.docx"
    word = Document()
    # 先对于文档英文为Arial，中文为宋体所有字体设置为小四，所有字体设置为小四，
    word.styles['Normal'].font.name = 'Arial'
    word.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    word.styles['Normal'].font.size = Pt(12)
    word.sections[0].left_margin = Cm(2.1)
    word.sections[0].right_margin = Cm(2.1)
    title = word.add_paragraph('湖南商务职业技术学院毕业设计答辩记录表')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.font.bold = True
    table = word.add_table(rows=7, cols=6, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头前三行对齐格式统一
    align_cccccl = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    align_clcl = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    row1 = []    #第1行的元素值用列表存储
    row1.append('学生姓名')
    row1.append(stu.last_name + stu.first_name)
    row1.append('二级学院')
    row1.append(stu.dept_name)
    row1.append('班级名称')
    row1.append(stu_project.class_name)
    for cell, align, ele in zip(table.rows[0].cells, align_cccccl, row1):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.alignment = align
        cell.paragraphs[0].add_run(ele)
    row2 = []  # 第2行的元素值用列表存储
    row2.append('专业名称')
    row2.append(stu.major)
    row2.append('学      号')
    row2.append(stu_no)
    row2.append('指导教师')
    row2.append(stu_project.instructor)
    for cell, align, ele in zip(table.rows[1].cells, align_cccccl, row2):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.alignment = align
        cell.paragraphs[0].add_run(ele)
    row3 = []  # 第3行的元素值用列表存储
    row3.append("选题名称")
    row3.append('')
    row3.append('')
    row3.append(stu_project.topic)
    row3.append("设计类型")
    row3.append(stu_defence.design_type)
    table.rows[2].cells[1].merge(table.rows[2].cells[3])
    for cell, align, ele in zip(table.rows[2].cells, (align_clcl*2)[:6], row3):
        cell.paragraphs[0].paragraph_format.alignment = align
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].add_run(ele)

    row4 = []  # 第4行的元素值用列表存储
    row4.append("答辩时间")
    row4.append('')
    row4.append('')
    row4.append(stu_defence.defence_date)
    row4.append("答辩地点")
    row4.append(stu_defence.defence_location)
    table.rows[3].cells[1].merge(table.rows[3].cells[3])
    for cell, align, ele in zip(table.rows[3].cells, (align_clcl * 2)[:6], row4):
        cell.paragraphs[0].paragraph_format.alignment = align
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].add_run(ele)

    row5 = []  # 第5行的元素值用列表存储
    row5.append("答辩组成员")
    row5.append('')
    row5.append('')
    row5.append('')
    row5.append('')
    inst_tmp = [stu_defence.def_inst1, stu_defence.def_inst2, stu_defence.def_inst3, stu_defence.def_inst4, stu_defence.def_inst5]
    inst_list = []
    inst_group = ''
    for i in inst_tmp:
        if i != 'nan':
            inst_list.append(i)
            inst_group = inst_group + i + "、"
    row5.append(inst_group.strip('、'))
    table.rows[4].cells[1].merge(table.rows[4].cells[5])
    for cell, align, ele in zip(table.rows[4].cells, (align_clcl * 2)[:6], row5):
        cell.paragraphs[0].paragraph_format.alignment = align
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].add_run(ele)
    for i in range(5):
        table.rows[i].height = Cm(1)

    table.rows[5].cells[0].merge(table.rows[5].cells[5])
    table.rows[5].cells[0].paragraphs[0].add_run('一、学生网络答辩影像（附：答辩截屏图片或录屏文件）')
    tmp6 = table.rows[5].cells[0].add_paragraph().add_run()
    try:
        tmp6.add_picture(f'media/{stu_project.defence_image}', height=Cm(7))
    except Exception:
        tmp6.add_picture(f'static/缺失.jpeg', width=Cm(1.8))
        tmp6.add_text("答辩影像未上传，请完善信息")
    table.rows[5].cells[0].add_paragraph()
    table.rows[5].cells[0].add_paragraph('二、学生自述内容')
    table.rows[5].cells[0].add_paragraph(stu_project.self_report)
    table.rows[5].cells[0].add_paragraph()
    table.rows[5].cells[0].add_paragraph('三、提问与问答')
    table.rows[5].cells[0].add_paragraph(stu_project.quiz.replace('\r',''))

    table.rows[6].cells[0].merge(table.rows[6].cells[2])
    table.rows[6].cells[3].merge(table.rows[6].cells[5])
    table.rows[6].cells[0].paragraphs[0].add_run('记录人（签字）：')
    tmp7 = table.rows[6].cells[0].add_paragraph().add_run()
    try:
        recorder = UserInfo.objects.get(last_name=stu_defence.recorder[0], first_name=stu_defence.recorder[1:])
        tmp7.add_picture(f'media/{recorder.sign}', width=Cm(1.8))
    except Exception:
        tmp7.add_picture(f'static/缺失.jpeg', width=Cm(1.8))
        tmp7.add_text(f"{recorder.last_name}{recorder.first_name}信息不完整，请联系其完善信息")
    table.rows[6].cells[0].paragraphs[1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[6].cells[0].add_paragraph(stu_defence.defence_date).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.rows[6].cells[3].paragraphs[0].add_run('答辩小组成员（签字）：')
    tmp72 = table.rows[6].cells[3].add_paragraph().add_run()
    for item in inst_list:
# 获取对应老师的UserInfo对象，然后获取相应的签名保存地址
        try:
            temp_inst = UserInfo.objects.get(last_name=item[0], first_name=item[1:])
            tmp72.add_picture(f'media/{temp_inst.sign}', width=Cm(1.8))
        except Exception:
            tmp72.add_picture(f'static/缺失.jpeg', width=Cm(1.8))
            tmp72.add_text(temp_inst.last_name + temp_inst.first_name + "信息不完整，请联系其完善信息")
    table.rows[6].cells[3].paragraphs[1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[6].cells[3].add_paragraph(stu_defence.defence_date).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    word.save(f'{path}/{word_name}')
    return word_name